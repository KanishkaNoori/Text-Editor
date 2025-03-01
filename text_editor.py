from tkinter import *
from tkinter import filedialog
from tkinter import font
from docx import Document
from tkinter import colorchooser


window = Tk()
window.title("EASY text!")
window.config(bg='#bfbfbf')

window.geometry("1200x600")
icon = PhotoImage(file="C:\\Users\\Kanishka\\Pictures\\application\\file.png")
window.iconphoto(True, icon)

global open_status_name
open_status_name = False

global selected
selected = False

def new_file():
    my_text.delete("1.0", END)
    window.title("New File - TextPad!")
    status_bar.config(text="New File        ")
    global open_status_name
    open_status_name = False
    
def open_file():
    my_text.delete("1.0", END)
    path = filedialog.askopenfilename(title='Open File', 
                                          filetypes=[("Text Files", "*.txt"),
                                                      ('Word Files', "*.docx"),

                                                      ("All Files", "*.*")])
    

    if path:
        a = Document(path)
        text = ""
        for i in a.paragraphs:            
            my_text.insert(END, i.text + "\n")
    # if path:
    #     if path.endswith(".docx"):
    #         a = Document(path)
    #         text=""
    #         for i in a.paragraphs:
    #             text += i.text + "\n"
    # else:          
    #     my_text.delete('1.0', END)
    #     text_file = filedialog.askopenfilename(initialdir="C:\\Users\\Kanishka\\Desktop" ,title="Open File",filetypes=(("Text Files", "*.txt"),("Word File","*.docx"),("HTML Files", "*.html"),("Python Files", "*.py"), ("All Files", "*.*")))

    
    # name = text_file
    # status_bar.config(text=f'{name}        ')
    # name = name.replace("C:\\Users\\Kanishka\\Desktop", "")
    # window.title(f'{name} - TextPad!')
    
    # path = open(path, 'r')
    # topic = path.read()
    # my_text.insert(END, topic)
    # path.close()
    
def save_as_file():
    text_file = filedialog.asksaveasfilename(defaultextension=".*", initialdir="C:\\Users\\Kanishka\\Desktop", title='Save File', filetypes=(("Text File","*.txt"),("Word File","*.docx"),("HTML files", "*.html"),("Python files", "*.py")
    ,("All files", "*.*"),
    ("document","*.docx")
    ))
    if text_file:
        te = my_text.get(1.0,END)
        if text_file.endswith(".docx"):
            doc = Document()
            doc.add_paragraph(te)
            doc.save(text_file)
        else:
            text_file = open(text_file,'w')
            text_file.write(my_text.get(1.0, END))
            text_file.close()
        name = text_file
        #status_bar.config(text=f'Saved: {name}        ')
        # name = name.replace("C:\\Users\\Kanishka\\Desktop", "")
        window.title(f'{name} - TextPad!')

        


def save_file():
    global open_status_name
    if open_status_name:
        text_file = open(open_status_name,'w')
        text_file.write(my_text.get(1.0, END))
        text_file.close()
        
        status_bar.config(text=f'Saved: {open_status_name}        ')
    else:
        save_as_file()

    
        
def cut_text(e):
    global selected
    if e:
        selected = window.clipboard_get()
    else:
        if my_text.selection_get():
            selected = my_text.selection_get()
            my_text.delete("sel.first", "sel.last")
            window.clipboard_clear()
            window.clipboard_append(selected)

def copy_text(e):
    global selected
    if e:
        selected = window.clipboard_get()
    if my_text.selection_get():
        selected = my_text.selection_get()
        window.clipboard_clear()
        window.clipboard_append(selected)

def paste_text(e):
    global selected
    if e:
        selected = window.clipboard_get()
    else:
        if selected:
            position = my_text.index(INSERT)
            my_text.insert(position, selected)
            
def bold_it():
    bold_font = font.Font(my_text, my_text.cget("font"))
    bold_font.configure(weight='bold')
    my_text.tag_configure('bold', font=bold_font)
    
    current_tags = my_text.tag_names("sel.first")
    
    if "bold" in current_tags:
        my_text.tag_remove("bold", "sel.first", "sel.last")
    else:
        my_text.tag_add("bold", "sel.first", "sel.last")
        
def italics_it():
    italics_font = font.Font(my_text, my_text.cget("font"))
    italics_font.configure(slant="italic")
    my_text.tag_configure('italic', font=italics_font)
    
    current_tags = my_text.tag_names("sel.first")
    
    if "italic" in current_tags:
        my_text.tag_remove("italic", "sel.first", "sel.last")
    else:
        my_text.tag_add("italic", "sel.first", "sel.last")

def text_color():
    my_color = colorchooser.askcolor()[1]
    if my_color:
        color_font = font.Font(my_text, my_text.cget("font"))
        my_text.tag_configure('colored', font=color_font, foreground=my_color)
        
        current_tags = my_text.tag_names("sel.first")
        
        if "colored" in current_tags:
            my_text.tag_remove("colored", "sel.first", "sel.last")
        else:
            my_text.tag_add("colored", "sel.first", "sel.last")
    
def bg_color():
    my_color = colorchooser.askcolor()[1]
    if my_color:
        my_text.config(bg=my_color)
    
def all_text_color():
    my_color = colorchooser.askcolor()[1]
    if my_color:
        my_text.config(fg=my_color)
 
def select_all(e):
    my_text.tag_add("sel", "1.0", END)

def clear_all(e):
    my_text.delete("1.0", END)
    
copy_image = PhotoImage(file="C:/Users/Kanishka/Pictures/copy.png")
cut_image = PhotoImage(file="C:/Users/Kanishka/Pictures/cut.png")
paste_image = PhotoImage(file="C:/Users/Kanishka/Pictures/paste.png")
undo_image = PhotoImage(file="C:/Users/Kanishka/Pictures/undo.png")
redo_image = PhotoImage(file="C:/Users/Kanishka/Pictures/redo.png")
select_image = PhotoImage(file="C:/Users/Kanishka/Pictures/select.png")
clear_image = PhotoImage(file="C:/Users/Kanishka/Pictures/clear_all.png")

new_image = PhotoImage(file="C:/Users/Kanishka/Pictures/new.png")
open_image = PhotoImage(file="C:/Users/Kanishka/Pictures/open.png")
copy_image = PhotoImage(file="C:/Users/Kanishka/Pictures/copy.png")
save_image = PhotoImage(file="C:/Users/Kanishka/Pictures/save.png")
save_as_image = PhotoImage(file="C:/Users/Kanishka/Pictures/save as.png")
exit_image = PhotoImage(file="C:/Users/Kanishka/Pictures/exit.png")

selected_image = PhotoImage(file="C:/Users/Kanishka/Pictures/selected.png")
all_text_image = PhotoImage(file="C:/Users/Kanishka/Pictures/text.png")
background_image = PhotoImage(file="C:/Users/Kanishka/Pictures/background.png")


toolbar_frame = Frame(window, bg='#bfbfbf')
toolbar_frame.pack(fill=X)

my_frame = Frame(window)
my_frame.pack(pady=5)

text_scroll = Scrollbar(my_frame)
text_scroll.pack(side=RIGHT, fill=Y)
text_scroll.config(width=20)

hor_scroll = Scrollbar(my_frame, orient='horizontal')
hor_scroll.pack(side=BOTTOM, fill=X)


my_text = Text(my_frame, width=97, height=25, font=("Helvetica", 16),padx = 40, pady=40 ,selectbackground='yellow', selectforeground="black", undo=True, yscrollcommand=text_scroll.set,xscrollcommand=hor_scroll.set, wrap="none")
my_text.pack()

text_scroll.config(command=my_text.yview)
hor_scroll.config(command=my_text.xview)

my_menu = Menu(window)
window.config(menu=my_menu)

file_menu = Menu(my_menu, tearoff=0, bg='#ffffff')
my_menu.add_cascade(label="File", menu=file_menu)
file_menu.add_command(label="New", command=new_file, image=new_image,compound=LEFT)
file_menu.add_command(label="Open", command=open_file, image=open_image,compound=LEFT)
file_menu.add_command(label="Save", command=save_file, image=save_image,compound=LEFT)
file_menu.add_command(label="Save As", command=save_as_file, image=save_as_image,compound=LEFT)
file_menu.add_separator()
file_menu.add_command(label="Exit", command=quit, image=exit_image,compound=LEFT)

edit_menu = Menu(my_menu, tearoff=0, bg='#ffffff')
my_menu.add_cascade(label="Edit", menu=edit_menu)
edit_menu.add_command(label="Cut", image=cut_image,compound=LEFT, command= lambda: cut_text(False), accelerator="(Ctrl+x)")
edit_menu.add_command(label="Copy",image=copy_image, compound=LEFT, command=lambda: copy_text(False), accelerator="(Ctrl+c)")
edit_menu.add_command(label="Paste ",image=paste_image, compound=LEFT, command=lambda: paste_text(False), accelerator="(Ctrl+v)")
edit_menu.add_separator()
edit_menu.add_command(label="Undo",image=undo_image, compound=LEFT, command=my_text.edit_undo, accelerator="(Ctrl+z)")
edit_menu.add_command(label="Redo",image=redo_image, compound=LEFT, command=my_text.edit_redo, accelerator="(Ctrl+y)")
edit_menu.add_separator()
edit_menu.add_command(label="Select All",image=select_image, compound=LEFT, command=lambda: select_all(True), accelerator="(Ctrl+a)")
edit_menu.add_command(label="Clear All",image=clear_image, compound=LEFT, command=lambda: clear_all(True), accelerator="(Ctrl+f)")

color_menu = Menu(my_menu, tearoff=0, bg='#ffffff')
my_menu.add_cascade(label="Colors", menu=color_menu)
color_menu.add_command(label="Selected Text", command= text_color, image=selected_image,compound=LEFT)
color_menu.add_command(label="All Text", command=all_text_color, image=all_text_image,compound=LEFT)
color_menu.add_command(label="Background",command=bg_color, image=background_image,compound=LEFT)

status_bar = Label(window, text='Ready        ', anchor=E)
status_bar.pack(fill=X, side=BOTTOM, ipady=5)

window.bind('<Control-Key-x>', cut_text)
window.bind('<Control-Key-c>', copy_text)
window.bind('<Control-Key-v>', paste_text)
window.bind('<Control-a>', select_all)
window.bind('<Control-f>', clear_all)

bold_button = Button(toolbar_frame, text='Bold', command=bold_it)
bold_button.grid(row=0, column=0, sticky=W, padx=5)

italics_button = Button(toolbar_frame, text='Italics', command=italics_it)
italics_button.grid(row=0, column=1, sticky=W, padx=5)

undo_button = Button(toolbar_frame, text='Undo', command=my_text.edit_undo)
undo_button.grid(row=0, column=2, sticky=W, padx=5)

redo_button = Button(toolbar_frame, text='Redo', command=my_text.edit_redo)
redo_button.grid(row=0, column=3, sticky=W, padx=5)

color_text_button = Button(toolbar_frame, text=' Text Color', command=text_color)
color_text_button.grid(row=0, column=4, padx=5)
window.mainloop()